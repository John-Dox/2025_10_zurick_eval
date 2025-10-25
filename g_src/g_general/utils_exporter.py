import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# --- FUNZIONE ESISTENTE (INVARIATA) ---
def export_to_word(session_log: list, base_dir: str):
    """
    Crea un documento Word formattato a partire dal log di una sessione di '5_ask.py'.
    """
    if not session_log:
        print("ℹ️  Nessuna interazione da salvare. Il file non verrà creato.")
        return

    try:
        doc = Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(11)

        title = doc.add_heading("Report Sessione di Lavoro RAG", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp_str = datetime.now().strftime("%d %B %Y, ore %H:%M:%S")
        subtitle = doc.add_paragraph(f"Generato il: {timestamp_str}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for i, turn_data in enumerate(session_log):
            doc.add_heading(f"Interazione #{i+1}", level=1)

            doc.add_heading("Domanda Utente", level=2)
            doc.add_paragraph(turn_data.get('question', 'N/D'))

            doc.add_heading("Percorso di Elaborazione", level=2)
            path_taken = turn_data.get('path_taken', 'sconosciuto')
            
            p = doc.add_paragraph()
            p.add_run("Tipo: ").bold = True
            if path_taken == 'structural_query':
                p.add_run("Query Strutturale")
                p_intent = doc.add_paragraph()
                p_intent.add_run("Intent Rilevato: ").bold = True
                p_intent.add_run(f"'{turn_data.get('analysis', {}).get('intent', 'N/D')}'")

            elif path_taken == 'follow_up':
                p.add_run("Follow-up su Interazione Precedente")
                p_style = doc.add_paragraph()
                p_style.add_run("Stile Richiesto: ").bold = True
                p_style.add_run(f"'{turn_data.get('analysis', {}).get('style', 'N/D')}'")
                p_orig_q = doc.add_paragraph()
                p_orig_q.add_run("Domanda Originale: ").bold = True
                p_orig_q.add_run(f"'{turn_data.get('original_question', 'N/D')}'")

            elif path_taken == 'rag':
                p.add_run("RAG (Retrieval-Augmented Generation)")
                if turn_data.get('context_chunks'):
                    doc.add_heading("Contesto Recuperato", level=3)
                    for chunk in turn_data['context_chunks']:
                        fonte_str = (
                            f"Fonte: [{chunk.payload.get('document_title', 'N/D')}] "
                            f"Art. {chunk.payload.get('articolo', 'N/A')}, "
                            f"Comma {chunk.payload.get('comma', 'N/A')} "
                            f"(Score: {chunk.score:.4f})"
                        )
                        doc.add_paragraph(fonte_str, style='Intense Quote')
                        doc.add_paragraph(chunk.payload.get('testo_originale_comma', ''))
                        doc.add_paragraph()

            elif path_taken == 'fallback':
                p.add_run("Fallback con Keyword")
            
            doc.add_heading("Risposta Finale", level=2)
            doc.add_paragraph(turn_data.get('final_answer', 'Nessuna risposta generata.'))
            
            if i < len(session_log) - 1:
                doc.add_page_break()

        reports_dir = os.path.join(base_dir, "e_reports", "01_words")
        os.makedirs(reports_dir, exist_ok=True)
        filename = f"Report_Sessione_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
        file_path = os.path.join(reports_dir, filename)
        doc.save(file_path)
        print(f"✅ Report della sessione salvato con successo in:\n   {file_path}")

    except Exception as e:
        print(f"❌ ERRORE CRITICO durante la creazione del file Word: {e}")

# --- NUOVA FUNZIONE PER L'EXPORT DELLA DIAGNOSTICA ---
def export_diagnostics_to_word(diagnostic_log: list, base_dir: str):
    """
    Crea un documento Word formattato a partire dal log di una sessione di diagnostica.
    """
    if not diagnostic_log:
        print("ℹ️  Nessuna diagnostica da salvare. Il file non verrà creato.")
        return

    try:
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        title = doc.add_heading("Report Diagnostica Keyword RAG", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp_str = datetime.now().strftime("%d %B %Y, ore %H:%M:%S")
        subtitle = doc.add_paragraph(f"Generato il: {timestamp_str}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for i, turn_data in enumerate(diagnostic_log):
            doc.add_heading(f"Diagnostica Query #{i+1}", level=1)
            
            p_query = doc.add_paragraph()
            p_query.add_run("Query Eseguita: ").bold = True
            p_query.add_run(f"'{turn_data.get('query', 'N/D')}'")
            doc.add_paragraph()

            doc.add_heading("Chunk Recuperati (Ordinati per Score)", level=2)
            retrieved_hits = turn_data.get('hits', [])

            if not retrieved_hits:
                doc.add_paragraph("Nessun chunk recuperato per questa query.")
            else:
                for j, hit in enumerate(retrieved_hits):
                    doc.add_heading(f"Risultato #{j+1}", level=3)
                    payload = hit.payload
                    
                    p_score = doc.add_paragraph()
                    p_score.add_run("Score: ").bold = True
                    p_score.add_run(f"{hit.score:.4f}")
                    p_score.add_run(f"\nFonte: ").bold = True
                    p_score.add_run(f"[{payload.get('document_title', 'N/D')}] Art. {payload.get('articolo')}, Comma {payload.get('comma')}")

                    p_text = doc.add_paragraph()
                    p_text.add_run("Testo: ").bold = True
                    p_text.add_run(f"\"{payload.get('testo_originale_comma', '')}\"")
                    
                    p_kw = doc.add_paragraph()
                    p_kw.add_run("Keywords:").bold = True
                    keywords_str = '\n'.join([f"- {kw}" for kw in payload.get('keywords', [])])
                    doc.add_paragraph(keywords_str)
                    doc.add_paragraph()

            if i < len(diagnostic_log) - 1:
                doc.add_page_break()

        reports_dir = os.path.join(base_dir, "e_reports")
        os.makedirs(reports_dir, exist_ok=True)
        filename = f"Report_Diagnostica_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
        file_path = os.path.join(reports_dir, filename)
        
        doc.save(file_path)
        print(f"\n✅ Report di diagnostica salvato con successo in:\n   {file_path}")

    except Exception as e:
        print(f"❌ ERRORE CRITICO durante la creazione del file Word di diagnostica: {e}")